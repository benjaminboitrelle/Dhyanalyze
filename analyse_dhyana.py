# -*- coding: utf-8 -*-

import os
import argparse
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches


def get_metadata(metadata):
    ''' Open metada file and store exposure time to corresponding images
    '''

    with open(metadata, 'r') as meta_file:
        file_content = meta_file.read().splitlines()

        # data looks like this: <exposurer_time>  <file_prefix>
        file_content = [s.split("\t") for s in file_content]
        for i, string in enumerate(file_content):
            try:
                string[0] = float(string[0])
            except ValueError:
                if string == ['']:
                    # remove empty lines
                    del file_content[i]
                else:
                    raise
    return sorted(file_content)


def get_meta(meta_name):
    ''' Return a matrix of the metada containing the images associated
        to an exposure time
    '''

    metadata = get_metadata(meta_name)
    meta = [[metadata[time][0],
            (metadata[time-1][1].strip(),
             metadata[time][1].strip())]
            for time, file in enumerate(metadata)
            if metadata[time-1][0] == metadata[time][0]]
    return meta


def get_images(meta, exposure_time, roi):
    ''' For a given exposure time return a stack of arrays corresponding
        to the 2 images taken
    '''
    img1 = np.array(plt.imread(meta[exposure_time][1][0] + ".tif"),
                    dtype=float)
    img2 = np.array(plt.imread(meta[exposure_time][1][1] + ".tif"),
                    dtype=float)
    img = np.dstack((img1[roi:-roi, roi:-roi],
                     img2[roi:-roi, roi:-roi]))

    return img


def plot_result(x, y, m, c, out_fname, title, label_x, label_y):
    ''' Show plot of data and data fitted
    '''
    plt.plot(x, y, '.', x, m*x+c, '-')
    plt.title(title)
    plt.xlabel(label_x)
    plt.ylabel(label_y)

    ax = plt.axes()
    if measurement_type == 'PTC':
        plt.text(0.8, 0.07,
                 'slope :  %.2f ADU/e- \n offset :  %.2f ADU' % (m, c),
                 fontsize=12,
                 horizontalalignment='center',
                 verticalalignment='center',
                 transform=ax.transAxes)

    elif measurement_type == 'Dark':
        plt.text(0.8, 0.07,
                 'slope :  %.2f ADU/s \n offset :  %.2f ADU' % (m*1000, c),
                 fontsize=12,
                 horizontalalignment='center',
                 verticalalignment='center',
                 transform=ax.transAxes)
    plt.savefig(out_fname)
    plt.show()
    plt.close()


def get_fit_parameters(x, y):
    ''' Linear fitting of data

        Output:
            slope, offset, residuals
    '''

    m, c = np.polyfit(x, y, 1)

    return m, c


if __name__ == "__main__":

    parser = argparse.ArgumentParser()
    parser.add_argument('--input',
                        dest='input_dir_path',
                        type=str,
                        help='Path to the input directory')

    parser.add_argument('--output',
                        dest='output_dir_path',
                        type=str,
                        help='Path to the output directory')

    parser.add_argument('--type',
                        dest='measurement_type',
                        type=str,
                        help='Type of measurements: \n Dark \n PTC')

    parser.add_argument('--roi',
                        dest='roi',
                        type=int,
                        nargs='?',
                        default=5,
                        help='Select a region of interest. This ROI is \
                              squared and centered. \n By default: \
                              10 pixels are excluded in vertically and \
                              horizontally.')

    args = parser.parse_args()
    input_dir_path = args.input_dir_path
    output_dir_path = args.output_dir_path
#    measurement_type = args.measurement_type
    roi = args.roi

    doc = Document()

# Start with characterisation of dark measurements
    measurement_type = 'Dark'
    os.chdir(input_dir_path + '/Dark')
    meta = get_meta('metadata.txt')
    time_range = len(meta)
    pic_dim = int(2048 - 2 * roi)
    print('picture_dimensions {}'.format(pic_dim))

# Create a stack of images depending on the exposure time:
#        img.shape() = (snap, nb_pixel_x, nb_pixel_y)
    img = np.asarray([get_images(meta, i, roi)
                      for i in range(time_range)])

    img_sum = np.sum(img, axis=3) / 2
    img_diff = (np.diff(img, axis=3) / np.sqrt(2)).reshape(time_range,
                                                           pic_dim,
                                                           pic_dim)
    avr_mat = np.mean(img_sum, axis=(1, 2))
    std_mat = np.std(img_diff, axis=(1, 2))
    exp_time = np.asarray([meta[time][0] for time in range(time_range)])
    slope_avr, offset_avr = get_fit_parameters(exp_time, avr_mat)
    slope_std, offset_std = get_fit_parameters(exp_time, std_mat)

# Calculate dark parameter and print plots
    dark_current = np.mean(avr_mat)
    min_std = np.min(std_mat)
    max_std = np.max(std_mat)
    dark_tmp_noise = np.mean(std_mat)

    print("Dark current {0:.2f} ADU".format(dark_current))
    print("Std min {0:.2f} ADU".format(min_std))
    print("Std max {0:.2f} ADU".format(max_std))
    print("Dark temporal noise {0:.2f} ADU".format(dark_tmp_noise))

    fig_mean_vs_timing = 'meanVsExposure_roi' + str(roi) + '.png'
    fig_std_vs_timing = 'standardDeviationVsExposure_roi' + str(roi) + '.png'
    os.chdir(output_dir_path)

    plot_result(exp_time,
                avr_mat,
                slope_avr,
                offset_avr,
                fig_mean_vs_timing,
                "Mean of the dark VS exposure time",
                "Exposure time [ms]",
                "Mean [ADU]")

    plot_result(exp_time,
                std_mat,
                slope_std,
                offset_std,
                fig_std_vs_timing,
                "Standard deviation of the dark VS exposure time",
                "Exposure time [ms]",
                "Standard deviation [ADU]")

# PTC
    measurement_type = 'PTC'
    input_dir_path_PTC = input_dir_path + '/PTC'

    input_dir_path_light = input_dir_path_PTC + "/Light"
    input_dir_path_dark = input_dir_path_PTC + "/Dark"

    os.chdir(input_dir_path_light)
    meta_light = get_meta("metadata.txt")
    time_range = len(meta_light)
    img_light = np.asarray([get_images(meta_light, snap, roi)
                            for snap in range(time_range)])

    os.chdir(input_dir_path_dark)
    meta_dark = get_meta("metadata.txt")
    time_range = len(meta_dark)
    img_dark = np.asarray([get_images(meta_dark, snap, roi)
                           for snap in range(time_range)])
    img_ptc = img_light - img_dark
    img_sum_ptc = np.sum(img_ptc, axis=3)/2
    img_diff_ptc = (np.diff(img_ptc, axis=3) / np.sqrt(2)).reshape(time_range,
                                                                   pic_dim,
                                                                   pic_dim)
    img_std_ptc = np.var(img_diff_ptc, axis=(1, 2), ddof=1)
    img_avr_ptc = np.mean(img_sum_ptc, axis=(1, 2))

    fit_bounds = np.where(np.logical_and(img_avr_ptc > 0,
                                         img_avr_ptc < 50000))
    slope, offset = get_fit_parameters(img_avr_ptc[fit_bounds],
                                       img_std_ptc[fit_bounds])

    saturation_adu = img_avr_ptc[img_std_ptc == np.max(img_std_ptc)]
    saturation_electron = saturation_adu / slope
    max_snr = np.log2(np.sqrt(saturation_electron))
    gain = dark_tmp_noise / slope

    print("Max of variance {0:.2f} ADU^2".format(np.max(img_std_ptc)))
    print("Saturation {0:.2f} ADU".format(saturation_adu[0]))
    print("Saturation {0:.2f} electrons".format(saturation_electron[0]))
    print("Max SNR {0:.2f}".format(max_snr[0]))
    print(slope, offset)

    os.chdir(output_dir_path)

    plot_result(img_avr_ptc,
                img_std_ptc,
                slope,
                offset,
                'stdVsMeanLight_roi'+str(roi)+'.png',
                'Transfer curve',
                'Mean value [ADU]',
                'Variance [ADU^2]')

    doc.add_heading('Low gain characeterisation', 0)
    doc.add_heading('Dark measurements', level=1)
    doc.add_paragraph('Size of ROI: {}x{} pixels'.format(roi, roi))
    doc.add_picture('meanVsExposure_roi'+str(roi)+'.png')
    doc.add_paragraph('Fitting results:')
    doc.add_paragraph('Dark current per second'
                      ' {0:.2f} ADU/s'.format(slope_avr*1000),
                      style='List Bullet')
    doc.add_paragraph('Dark current {0:.2f} ADU'.format(dark_current),
                      style='List Bullet')
    doc.add_paragraph('Below is shown the standard deviation of the dark'
                      ' (in ADU) as a function of the exposure time'
                      ' (in milliseconds)')
    doc.add_picture('standardDeviationVsExposure_roi'+str(roi)+'.png')
    doc.add_paragraph('Dark temporal noise'
                      ' {0:.2f} ADU'.format(dark_tmp_noise),
                      style='List Bullet')
    doc.add_paragraph('Standard deviation min {0:.2f} ADU'.format(min_std),
                      style='List Bullet')
    doc.add_paragraph('Standard deviation max {0:.2f} ADU'.format(max_std),
                      style='List Bullet')

    doc.add_heading('PTC measurements', level=1)
    doc.add_picture('stdVsMeanLight_roi'+str(roi)+'.png')
    doc.add_paragraph('Results of PTC measurement:')
    doc.add_paragraph('Gain {0:.2f} ADU/e-'.format(slope),
                      style='List Bullet')
    doc.add_paragraph('Saturation {0:.2f} ADU'.format(saturation_adu[0]),
                      style='List Bullet')
    doc.add_paragraph('Saturation {0:.2f} e-'.format(saturation_electron[0]),
                      style='List Bullet')
    doc.add_paragraph('Max SNR {0:.2f} bit'.format(max_snr[0]),
                      style='List Bullet')
    doc.add_paragraph('Temporal dark noise {0:.2f} e-'.format(gain),
                      style='List Bullet')
    table = doc.add_table(rows=4, cols=2)
    cell = table.cell(0, 0)
    cell.text = 'Parameter'
    cell = table.cell(0, 1)
    cell.text = 'Low Gain'
    cell = table.cell(1, 0)
    cell.text = 'Overall gain (DN/e-)'
    cell = table.cell(1, 1)
    cell.text = '{}'.format(gain)
    cell = table.cell(2, 0)
    cell.text = 'Saturation (e-)'
    cell = table.cell(2, 1)
    cell.text = '{}'.format(saturation_electron[0])
    cell = table.cell(3, 0)
    cell.text = 'Dark current (ADU)'
    cell = table.cell(3, 1)
    cell.text = '{}'.format(dark_current)

    doc.save('lowGain_roi'+str(roi)+'.docx')
